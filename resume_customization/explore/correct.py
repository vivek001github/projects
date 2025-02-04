import streamlit as st
import google.generativeai as genai
import os
import pdfplumber
from docx import Document
from docx.shared import Pt
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
import re

# Configure the API key for Google Generative AI
genai.configure(api_key=os.environ["API_KEY"])

# Function to read a DOCX file
def read_docx(file):
    doc = Document(file)
    return '\n'.join([para.text for para in doc.paragraphs])

# Function to read a PDF file using pdfplumber
def read_pdf(file):
    text = ""
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text

# Function to customize the resume based on the job description
def customize_resume_api(resume_text, job_description):
    prompt = f"""
    Customize the following resume to match the provided job description while retaining all details from the original resume. 

    Make sure to highlight relevant skills, experiences, and qualifications from the job description, while also including all sections of the original resume. 

    Structure: Organize the resume into clear sections, including "Profile," "Education," "Skills & Abilities," "Projects & Experience," and "Internships." Add a "Declaration" section at the end.
    Formatting:
    - Display the candidate's name in large, bold black text at the top, with contact details directly below.
    - Use bold black text for all section headings and subtitles, ignoring any colors from the original resume. All text should be black, without additional colors or styling.
    - Use minimal and consistent spacing above and below each section title and subtitle, creating a neat, compact layout without extra gaps.
    
    Content Presentation:
    - Retain all details from the original resume while integrating relevant information from the job description.
    - In the "Skills & Abilities" section, include the following subsections formatted as comma-separated lists:
      - Technical Skills: List technical skills relevant to the job description, separated by commas.
      - Software Skills: List software and tools that the candidate is proficient in, separated by commas.
      - Soft Skills: List soft skills that highlight interpersonal and leadership abilities, separated by commas.
    - Ensure that the "Skills & Abilities" section contains no explanations or extra descriptions; it should only list the skills under the specified categories in the order: Technical Skills, Software Skills, and Soft Skills.
    - Align all sections consistently, with uniform line spacing of 0 lines, and avoid any extra indentation or spaces.
    
    Final Appearance: Ensure the resume is professional, with black text only and a compact layout. Use consistent line spacing of 0 lines and minimal spacing between titles, subtitles, and content.
    Important: Provide only the customized resume content without any additional notes, explanations, or formatting instructions.

    Resume: {resume_text}

    Job Description: {job_description}
    """
    try:
        model = genai.GenerativeModel("gemini-1.5-flash")
        response = model.generate_content(prompt)
        
        # Clean any markdown formatting like **bold** and *bullets* using regex
        clean_text = re.sub(r'\*{1,2}(.*?)\*{1,2}', r'\1', response.text)  # Remove any bold or bullet markers
        return clean_text.strip()
    except Exception as e:
        st.error(f"Error customizing resume: {e}")
        return None

# Function to save the customized resume to a DOCX file with formatting
def save_to_docx(customized_resume):
    doc = Document()

    # Set custom margins
    sections = doc.sections
    for section in sections:
        section.left_margin = Pt(36)  # 0.5 inch margin
        section.right_margin = Pt(36)
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
    
    # Define titles and specific lines for custom formatting
    name_added = False
    title_set = {"Profile", "Education", "Skills & Abilities", "Projects & Experience", "Internships", "Declaration"}
    project_titles = [
        "AUTOMATIC PAPER CUTTING MACHINE USING GENEVA MECHANISM",
        "SOCIAL DISTANCING ID CARD",
        "Facial Landmark Detection",
        "QUIZ GENERATOR"
    ]
    education_titles = {"Bachelor of Engineering", "Pre University Education", "C.B.S.C"}

    lines = customized_resume.split("\n")
    
    for line in lines:
        line = line.strip()  # Remove leading and trailing whitespace
        if not line:
            continue
        
        if not name_added:
            # First non-empty line is assumed to be the name
            name_para = doc.add_paragraph()
            name_run = name_para.add_run(line)
            name_run.font.size = Pt(20)
            name_run.bold = True
            name_para.paragraph_format.space_after = Pt(0)  # Remove space after name paragraph
            name_added = True
        elif line in title_set:
            # Section titles like Profile, Education, etc.
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(line)
            title_run.font.size = Pt(14)
            title_run.bold = True
            title_para.paragraph_format.space_before = Pt(12)
            title_para.paragraph_format.space_after = Pt(0)
        elif line in project_titles:
            # Bold project titles in Projects & Experience section
            project_para = doc.add_paragraph()
            project_run = project_para.add_run(line)
            project_run.font.size = Pt(11)
            project_run.bold = True
            project_para.paragraph_format.space_before = Pt(12)  # Add space before each project title
            project_para.paragraph_format.space_after = Pt(0)
        elif line in education_titles:
            # Bold and size 11 for specific education lines
            education_para = doc.add_paragraph()
            education_run = education_para.add_run(line)
            education_run.font.size = Pt(11)
            education_run.bold = True
            education_para.paragraph_format.space_after = Pt(0)
        elif line in ["CGPA: 7 (upto 5th semester)", "Performance: 76%"]:
            # Add space after specific lines for CGPA and Performance
            specific_para = doc.add_paragraph()
            specific_run = specific_para.add_run(line)
            specific_run.font.size = Pt(11)
            specific_para.paragraph_format.space_after = Pt(12)  # Add space after these lines
        else:
            # Regular content
            content_para = doc.add_paragraph()
            content_run = content_para.add_run(line)
            content_run.font.size = Pt(11)
            content_para.paragraph_format.space_after = Pt(0)
            content_para.paragraph_format.line_spacing = Pt(12)

    # Save to a BytesIO stream
    byte_stream = BytesIO()
    doc.save(byte_stream)
    byte_stream.seek(0)
    return byte_stream

# Function to save the customized resume to a PDF file
def save_to_pdf(customized_resume):
    byte_stream = BytesIO()
    pdf = canvas.Canvas(byte_stream, pagesize=letter)
    pdf.setFont("Helvetica", 12)
    
    left_margin = 0.5 * inch
    top_margin = 11 * inch - 0.5 * inch
    
    text = pdf.beginText(left_margin, top_margin)
    text.setFont("Helvetica", 12)
    text.setLeading(12)  # Set line spacing to be consistent and minimal
    
    for line in customized_resume.split('\n'):
        line = line.strip()  # Remove leading and trailing whitespace
        if line:  # Only add non-empty lines
            text.textLine(line)
    
    pdf.drawText(text)
    pdf.showPage()
    pdf.save()
    
    byte_stream.seek(0)
    return byte_stream

# Streamlit app layout
def main():
    st.title("Resume Customization Tool")

    uploaded_resume = st.file_uploader("Upload your base resume (in .txt, .docx, or .pdf format)", type=["txt", "docx", "pdf"])
    uploaded_job_description = st.file_uploader("Upload the job description (in .txt format only)", type=["txt"])
    job_description_text = st.text_area("Or, enter the job description manually", height=150)

    if uploaded_resume:
        if uploaded_resume.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            resume_text = read_docx(uploaded_resume)
        elif uploaded_resume.type == "application/pdf":
            resume_text = read_pdf(uploaded_resume)
        else:
            resume_text = uploaded_resume.read().decode("utf-8")
        
        st.subheader("Uploaded Resume Content")
        st.text_area("Base Resume", resume_text, height=300)

        if uploaded_job_description:
            job_description = uploaded_job_description.read().decode("utf-8")
        else:
            job_description = job_description_text

        st.subheader("Job Description Content")
        st.text_area("Job Description", job_description, height=300)

        if st.button("Customize Resume"):
            customized_resume = customize_resume_api(resume_text, job_description)
            if customized_resume:
                st.subheader("Customized Resume")
                st.text_area("Your Customized Resume", customized_resume, height=400)

                docx_file = save_to_docx(customized_resume)
                pdf_file = save_to_pdf(customized_resume)

                st.download_button(
                    label="Download Customized Resume as DOCX",
                    data=docx_file,
                    file_name="customized_resume.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

                st.download_button(
                    label="Download Customized Resume as PDF",
                    data=pdf_file,
                    file_name="customized_resume.pdf",
                    mime="application/pdf"
                )

if __name__ == "__main__":
    main()
